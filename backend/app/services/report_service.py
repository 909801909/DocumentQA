from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from docx import Document as DocxDocument
from docx.shared import Inches
import base64
from typing import List, Dict, Optional
import os

# 字体查找和注册
def register_chinese_font():
    font_name = "SimSun"
    # 在常见路径中查找字体文件
    font_paths = [
        'C:/Windows/Fonts/simsun.ttc',  # Windows
        '/usr/share/fonts/truetype/simsun/simsun.ttc', # Linux
        'simsun.ttc' # 项目根目录
    ]
    
    font_path = None
    for path in font_paths:
        if os.path.exists(path):
            font_path = path
            break
            
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont(font_name, font_path))
            return font_name
        except Exception as e:
            print(f"Warning: Found font at {font_path} but failed to register: {e}")
    
    print("Warning: No suitable Chinese font (simsun.ttc) found. PDF reports may not display Chinese characters correctly.")
    return "Helvetica" # Fallback font

CHINESE_FONT = register_chinese_font()

class ReportService:
    def generate_report(self, format: str, content: Dict, selected_docs: List[Dict], kg_image_base64: Optional[str]) -> bytes:
        if format.lower() == 'pdf':
            return self._generate_pdf(content, selected_docs, kg_image_base64)
        elif format.lower() == 'word':
            return self._generate_word(content, selected_docs, kg_image_base64)
        else:
            raise ValueError("Unsupported report format")

    def _generate_pdf(self, content: Dict, selected_docs: List[Dict], kg_image_base64: Optional[str]) -> bytes:
        buffer = BytesIO()
        styles = getSampleStyleSheet()
        # 应用注册好的中文字体
        styles['Title'].fontName = CHINESE_FONT
        styles['h1'].fontName = CHINESE_FONT
        styles['h2'].fontName = CHINESE_FONT
        styles['Normal'].fontName = CHINESE_FONT
        
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []

        story.append(Paragraph(content.get('title', '分析报告'), styles['Title']))
        story.append(Spacer(1, 12))

        story.append(Paragraph('1. 报告摘要', styles['h1']))
        story.append(Paragraph(content.get('summary', ''), styles['Normal']))
        story.append(Spacer(1, 12))

        story.append(Paragraph('2. 分析文档列表', styles['h1']))
        for doc_item in selected_docs:
            story.append(Paragraph(f"- {doc_item['filename']}", styles['Normal']))
        story.append(Spacer(1, 12))

        story.append(Paragraph('3. 知识图谱分析', styles['h1']))
        if kg_image_base64:
            img_data = base64.b64decode(kg_image_base64)
            # 修正：为 reportlab 使用 inch 单位
            img = Image(BytesIO(img_data), width=6*inch, height=4.5*inch)
            img.hAlign = 'CENTER'
            story.append(img)
        else:
            story.append(Paragraph("未能生成知识图谱图像。", styles['Normal']))

        doc.build(story)
        return buffer.getvalue()

    def _generate_word(self, content: Dict, selected_docs: List[Dict], kg_image_base64: Optional[str]) -> bytes:
        doc = DocxDocument()
        doc.add_heading(content.get('title', '分析报告'), 0)

        doc.add_heading('1. 报告摘要', level=1)
        doc.add_paragraph(content.get('summary', ''))

        doc.add_heading('2. 分析文档列表', level=1)
        for doc_item in selected_docs:
            doc.add_paragraph(doc_item['filename'], style='List Bullet')

        doc.add_heading('3. 知识图谱分析', level=1)
        if kg_image_base64:
            img_data = base64.b64decode(kg_image_base64)
            # Word 文档使用 Inches 单位
            doc.add_picture(BytesIO(img_data), width=Inches(6.0))
        else:
            doc.add_paragraph("未能生成知识图谱图像。")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
